import asyncio
import io
import json
import logging
import tempfile
import time
import os
from typing import List, Dict, Any, Optional
from io import BytesIO
import zipfile

from aiogram import Bot
from aiogram.types import Message
from langchain_community.document_loaders import PyPDFLoader
from langchain.schema import Document
from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract
from docx import Document as DocxDocument
from openpyxl import load_workbook

from ..config import Settings
from ..storage.vector import VectorStore
from ..text.chunking import Chunker


class DocumentIngestionService:
    """Сервис по обработке и индексации документов."""

    def __init__(self, settings: Settings, vector_store: VectorStore):
        """Создает сервис с настройками приложения."""

        self._settings = settings
        self._vector_store = vector_store
        self._logger = logging.getLogger(__name__)
        self._chunker = Chunker(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
        )
        # Детектор наличия poppler (pdfinfo) для OCR PDF
        self._pdfinfo_available = self._detect_pdfinfo()
        # Калибровочные данные о производительности моделей эмбеддингов
        self._calibration_file = "/app/data/embedding_calibration.json"
        self._calibration_data: Dict[str, float] = self._load_calibration()
        # Загружаем пресеты моделей
        self._model_presets: Dict[str, Dict[str, Any]] = self._load_model_presets()

    # ----- Публичные операции над хранилищем (SRP/инкапсуляция VectorStore) -----
    def clear_store(self) -> None:
        """Очищает векторное хранилище документов."""

        self._vector_store.clear()

    def document_count(self) -> int:
        """Возвращает количество документов в хранилище."""

        return self._vector_store.document_count()

    def file_count(self) -> int:
        """Возвращает количество файлов-источников по метаданным."""

        return self._vector_store.file_count()

    def _load_model_presets(self) -> Dict[str, Dict[str, Any]]:
        """Загружает информацию о пресетах моделей эмбеддингов."""
        try:
            preset_path = os.environ.get("EMBED_PRESETS_JSON", "/app/web/embed_presets.json")
            with open(preset_path, "r", encoding="utf-8") as f:
                presets = json.load(f)
            result = {}
            for preset in presets:
                name = preset.get("name") or preset.get("pull", "")
                if name:
                    result[name] = preset
            return result
        except Exception:
            return {}

    def _load_calibration(self) -> Dict[str, float]:
        """Загружает калибровочные данные о скорости обработки моделей (чанков в секунду)."""
        try:
            if os.path.exists(self._calibration_file):
                with open(self._calibration_file, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception:
            pass
        return {}

    def _save_calibration(self) -> None:
        """Сохраняет калибровочные данные на диск."""
        try:
            os.makedirs(os.path.dirname(self._calibration_file), exist_ok=True)
            with open(self._calibration_file, "w", encoding="utf-8") as f:
                json.dump(self._calibration_data, f, indent=2)
        except Exception as e:
            self._logger.warning(f"Не удалось сохранить калибровку: {e}")

    def _calibrate_model(self, model_name: str) -> float:
        """Выполняет калибровку модели эмбеддингов и возвращает скорость (чанков в секунду)."""
        try:
            test_text = "Это тестовый текст для калибровки производительности модели эмбеддингов." * 20
            test_chunks = [test_text[:self._settings.chunk_size] for _ in range(3)]
            
            start_time = time.time()
            for chunk in test_chunks:
                try:
                    self._vector_store._embeddings.embed_query(chunk)
                except Exception:
                    pass
            elapsed = time.time() - start_time
            
            if elapsed > 0:
                chunks_per_second = len(test_chunks) / elapsed
                self._calibration_data[model_name] = chunks_per_second
                self._save_calibration()
                self._logger.info(f"Калибровка модели {model_name}: {chunks_per_second:.2f} чанков/сек")
                return chunks_per_second
        except Exception as e:
            self._logger.warning(f"Ошибка калибровки модели {model_name}: {e}")
        
        return self._estimate_chunks_per_second_by_model_size(model_name)

    def _estimate_chunks_per_second_by_model_size(self, model_name: str) -> float:
        """Оценивает скорость на основе размера модели из пресетов."""
        preset = self._model_presets.get(model_name, {})
        vram = preset.get("vram", "")
        
        try:
            vram_mb = int(vram) if vram != "—" else 1000
            
            if vram_mb < 1000:
                return 8.0
            elif vram_mb < 5000:
                return 4.0
            elif vram_mb < 10000:
                return 2.5
            elif vram_mb < 20000:
                return 1.5
            else:
                return 0.8
        except Exception:
            return 2.0

    def _get_chunks_per_second(self, model_name: str) -> float:
        """Возвращает скорость обработки чанков для модели (с калибровкой при необходимости)."""
        if model_name in self._calibration_data:
            return self._calibration_data[model_name]
        
        return self._estimate_chunks_per_second_by_model_size(model_name)

    def calibrate_current_model_background(self) -> None:
        """Запускает калибровку текущей модели в фоновом режиме."""
        model_name = self._settings.embeddings_model
        if not model_name:
            return
        
        if model_name in self._calibration_data:
            return
        
        try:
            self._logger.info(f"Запуск фоновой калибровки модели {model_name}")
            self._calibrate_model(model_name)
        except Exception as e:
            self._logger.warning(f"Ошибка фоновой калибровки: {e}")

    def _estimate_text_chars_from_file(self, file_size_bytes: int, filename: str) -> int:
        """Оценивает количество символов текста в файле на основе его типа и размера."""
        suffix = filename.lower()
        
        if suffix.endswith(".pdf"):
            return int(file_size_bytes * 0.5)
        elif suffix.endswith(".docx"):
            return int(file_size_bytes * 0.4)
        elif suffix.endswith(".xlsx"):
            return int(file_size_bytes * 0.3)
        elif any(suffix.endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp")):
            return int(file_size_bytes * 0.02)
        elif suffix.endswith(".json"):
            return int(file_size_bytes * 0.95)
        else:
            return int(file_size_bytes * 0.8)

    def estimate_processing_time(self, file_size_bytes: int, filename: str) -> str:
        """Оценивает примерное время обработки файла с учетом модели эмбеддингов и реальной производительности."""
        
        estimated_chars = self._estimate_text_chars_from_file(file_size_bytes, filename)
        
        chunk_size = self._settings.chunk_size
        estimated_chunks = max(1, estimated_chars // chunk_size)
        
        model_name = self._settings.embeddings_model or "unknown"
        chunks_per_second = self._get_chunks_per_second(model_name)
        
        embedding_time = estimated_chunks / chunks_per_second if chunks_per_second > 0 else estimated_chunks * 0.5
        
        overhead_time = 10.0
        if filename.lower().endswith(".pdf"):
            # PDF требует больше времени на парсинг
            overhead_time += max(20.0, file_size_bytes / (1024 * 1024) * 2)
        elif any(filename.lower().endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp")):
            overhead_time += 10.0
        
        total_seconds = embedding_time + overhead_time
        
        total_seconds = max(total_seconds, 5)
        
        if total_seconds < 60:
            return f"{int(total_seconds)} сек"
        else:
            minutes = int(total_seconds // 60)
            seconds = int(total_seconds % 60)
            if seconds > 0:
                return f"{minutes} мин {seconds} сек"
            else:
                return f"{minutes} мин"

    async def _add_documents_with_progress(
        self, 
        documents: List[Document], 
        message: Message, 
        bot: Bot
    ) -> None:
        """Добавляет документы в векторную БД с отображением прогресса."""
        
        total = len(documents)
        if total == 0:
            return
        
        # Отправляем начальное сообщение о прогрессе
        progress_msg = await message.answer(
            text=f"Векторизация: 0/{total} чанков (0%)\nОсталось: расчет..."
        )
        
        start_time = time.time()
        last_update_time = start_time
        
        # Обрабатываем документы небольшими батчами для обновления прогресса
        batch_size = 10
        processed = 0
        
        for i in range(0, total, batch_size):
            batch = documents[i:i + batch_size]
            
            # Добавляем батч в векторную БД
            await self._vector_store.add_documents(batch)
            processed += len(batch)
            
            # Обновляем прогресс каждые 5 секунд
            current_time = time.time()
            if current_time - last_update_time >= 5.0 or processed >= total:
                elapsed = current_time - start_time
                progress_pct = int(processed / total * 100)
                
                # Оценка оставшегося времени
                if processed > 0 and elapsed > 0:
                    speed = processed / elapsed
                    remaining = total - processed
                    eta_seconds = remaining / speed if speed > 0 else 0
                    
                    if eta_seconds < 60:
                        eta_str = f"{int(eta_seconds)} сек"
                    else:
                        eta_minutes = int(eta_seconds // 60)
                        eta_secs = int(eta_seconds % 60)
                        if eta_secs > 0:
                            eta_str = f"{eta_minutes} мин {eta_secs} сек"
                        else:
                            eta_str = f"{eta_minutes} мин"
                else:
                    eta_str = "расчет..."
                
                try:
                    await bot.edit_message_text(
                        chat_id=progress_msg.chat.id,
                        message_id=progress_msg.message_id,
                        text=f"Векторизация: {processed}/{total} чанков ({progress_pct}%)\nОсталось: {eta_str}"
                    )
                    last_update_time = current_time
                except Exception:
                    pass
        
        # Удаляем сообщение о прогрессе
        try:
            await bot.delete_message(chat_id=progress_msg.chat.id, message_id=progress_msg.message_id)
        except Exception:
            pass

    async def ingest_from_message(self, message: Message, bot: Bot) -> None:
        """Загружает документ из сообщения и индексирует его."""

        status_message = None
        try:
            buffer = io.BytesIO()
            await bot.download(message.document, buffer)
            buffer.seek(0)
            filename = message.document.file_name
            file_bytes = buffer.read()
            self._logger.info("Получен файл: %s, размер байт: %s", filename, len(file_bytes))
            
            documents = await self._load_documents(
                filename=filename,
                file_bytes=file_bytes,
            )
            self._logger.info("Загружено документов из файла: %s", len(documents))
            
            chunks = self._chunker.chunk_documents(documents=documents)
            self._logger.info("Сформировано чанков: %s", len(chunks))
            
            # Фильтруем пустые чанки во избежание пустых эмбеддингов
            non_empty = [d for d in chunks if d.page_content and d.page_content.strip()]
            if not non_empty:
                await message.answer(text="Не удалось извлечь текст из документа или он пуст.")
                self._logger.warning("Не удалось извлечь текст: %s", filename)
                return
            
            # Векторизация с прогресс-баром
            await self._add_documents_with_progress(non_empty, message, bot)
            self._logger.info("Добавлено чанков в векторную БД: %s", len(non_empty))
        except Exception:
            self._logger.exception("Критическая ошибка при обработке файла: %s", getattr(message.document, 'file_name', 'unknown'))
            await message.answer(text="Ошибка при обработке файла. Попробуйте позже.")

    async def ingest_photo_from_message(self, message: Message, bot: Bot) -> None:
        """Загружает фотографию из сообщения (photo) и индексирует её через OCR."""

        if not message.photo:
            return
        
        try:
            buffer = io.BytesIO()
            # Берем самый большой доступный размер
            try:
                await bot.download(message.photo[-1], buffer)
            except Exception:
                self._logger.exception("Не удалось скачать фото")
                await message.answer(text="Ошибка при загрузке фото. Попробуйте позже.")
                return
            buffer.seek(0)
            # Синтетическое имя для маршрутизации обработчика как изображения
            filename = f"photo_{getattr(message, 'message_id', 'unknown')}.jpg"
            file_bytes = buffer.read()
            self._logger.info("Получено фото: %s байт", len(file_bytes))
            
            documents = await self._load_documents(
                filename=filename,
                file_bytes=file_bytes,
            )
            self._logger.info("Загружено документов из фото: %s", len(documents))
            
            chunks = self._chunker.chunk_documents(documents=documents)
            self._logger.info("Сформировано чанков (photo): %s", len(chunks))
            
            non_empty = [d for d in chunks if d.page_content and d.page_content.strip()]
            if not non_empty:
                await message.answer(text="Не удалось извлечь текст с изображения.")
                self._logger.warning("Не удалось извлечь текст: photo %s", getattr(message, "message_id", 0))
                return
            
            # Векторизация с прогресс-баром
            await self._add_documents_with_progress(non_empty, message, bot)
            self._logger.info("Добавлено чанков из фото в векторную БД: %s", len(non_empty))
        except Exception:
            self._logger.exception("Критическая ошибка при обработке фото")
            await message.answer(text="Ошибка при обработке фото. Попробуйте позже.")

    async def _load_documents(self, filename: str, file_bytes: bytes) -> List[Document]:
        """Загружает документы в зависимости от формата файла."""

        suffix = filename.lower()
        if suffix.endswith(".pdf"):
            self._logger.info("Начало загрузки PDF документа: %s", filename)
            docs = await self._load_pdf(file_bytes=file_bytes)
            self._logger.info("PDF загружен, страниц: %s", len(docs))
            # Ограничение на размер файла для OCR: не более 10 МБ
            if len(file_bytes) <= 10 * 1024 * 1024:
                self._logger.info("Запуск OCR для PDF")
                ocr_docs = await self._ocr_pdf(file_bytes)
                self._logger.info("OCR завершен, получено документов: %s", len(ocr_docs))
                return docs + ocr_docs
            else:
                self._logger.info("Файл слишком большой для OCR (%s байт), пропускаем OCR", len(file_bytes))
                return docs
        if suffix.endswith(".json"):
            return self._load_json(file_bytes=file_bytes, filename=filename)
        if suffix.endswith(".docx"):
            return await self._load_docx(file_bytes=file_bytes, filename=filename)
        if suffix.endswith(".xlsx"):
            return await self._load_xlsx(file_bytes=file_bytes, filename=filename)
        if any(suffix.endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp")):
            return [self._load_image_ocr(file_bytes=file_bytes, filename=filename)]
        content = file_bytes.decode("utf-8", errors="ignore")
        return [Document(page_content=content, metadata={"source": filename})]

    async def _load_pdf(self, file_bytes: bytes) -> List[Document]:
        """Читает PDF документ в асинхронном контексте."""

        loop = asyncio.get_running_loop()
        with tempfile.NamedTemporaryFile(suffix=".pdf") as temp_file:
            temp_file.write(file_bytes)
            temp_file.flush()
            loader = PyPDFLoader(temp_file.name)
            documents = await loop.run_in_executor(None, loader.load)
        return documents

    def _load_json(self, file_bytes: bytes, filename: str) -> List[Document]:
        """Читает JSON документ и формирует список документов."""

        data = json.loads(file_bytes.decode("utf-8"))
        if isinstance(data, dict):
            data = [data]
        documents = []
        for index, item in enumerate(data):
            content = json.dumps(item, ensure_ascii=False)
            documents.append(
                Document(
                    page_content=content,
                    metadata={"source": filename, "index": index},
                )
            )
        return documents

    async def _load_docx(self, file_bytes: bytes, filename: str) -> List[Document]:
        """Извлекает текст и изображения (OCR) из DOCX."""

        with tempfile.NamedTemporaryFile(suffix=".docx") as tf:
            tf.write(file_bytes)
            tf.flush()
            docx = DocxDocument(tf.name)
            text_parts: List[str] = [p.text for p in docx.paragraphs if p.text and p.text.strip()]

        ocr_texts: List[str] = []
        with zipfile.ZipFile(BytesIO(file_bytes)) as zf:
            for name in zf.namelist():
                if name.startswith("word/media/") and name.lower().split(".")[-1] in {"png", "jpg", "jpeg", "webp", "tif", "tiff", "bmp"}:
                    try:
                        img_bytes = zf.read(name)
                        img = Image.open(BytesIO(img_bytes))
                        ocr_text = pytesseract.image_to_string(img)
                        if ocr_text and ocr_text.strip():
                            ocr_texts.append(ocr_text)
                    except Exception:
                        self._logger.exception("OCR DOCX image failed: %s", name)

        joined = "\n".join(text_parts + ocr_texts)
        return [Document(page_content=joined, metadata={"source": filename})]

    async def _load_xlsx(self, file_bytes: bytes, filename: str) -> List[Document]:
        """Извлекает текст из ячеек XLSX."""

        with tempfile.NamedTemporaryFile(suffix=".xlsx") as tf:
            tf.write(file_bytes)
            tf.flush()
            wb = load_workbook(tf.name, data_only=True, read_only=True)
            texts: List[str] = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    for value in row:
                        if isinstance(value, str) and value.strip():
                            texts.append(value)
        joined = "\n".join(texts)
        return [Document(page_content=joined, metadata={"source": filename})]

    def _load_image_ocr(self, file_bytes: bytes, filename: str) -> Document:
        """Прогон OCR по отдельному изображению."""

        img = Image.open(BytesIO(file_bytes))
        text = pytesseract.image_to_string(img)
        return Document(page_content=text or "", metadata={"source": filename, "ocr": True})

    async def _ocr_pdf(self, file_bytes: bytes) -> List[Document]:
        """OCR по страницам PDF и возврат объединенных текстов."""

        docs: List[Document] = []
        if not self._pdfinfo_available:
            # Без poppler (pdfinfo) pdf2image не сможет работать. Пропускаем бесшумно.
            return docs
        try:
            images = convert_from_bytes(file_bytes, dpi=200)
            ocr_texts: List[str] = []
            for index, img in enumerate(images):
                txt = pytesseract.image_to_string(img)
                if txt and txt.strip():
                    ocr_texts.append(f"[Страница {index+1}]\n{txt}")
            if ocr_texts:
                docs.append(Document(page_content="\n\n".join(ocr_texts), metadata={"source": "pdf-ocr"}))
        except Exception:
            self._logger.exception("OCR PDF failed")
        return docs

    def _detect_pdfinfo(self) -> bool:
        """Проверяет наличие утилиты pdfinfo (Poppler) в PATH."""
        try:
            import shutil

            return shutil.which("pdfinfo") is not None
        except Exception:
            return False

